import fitz
import base64
import io
from docx import Document
from typing import List
from services.llm_client import describe_image

async def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyMuPDF and handle images vision-wise."""
    text = ""
    try:
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            
            # 1. 提取普通文字
            text += page.get_text() + "\n"
            
            # 2. 提取图片并进行视觉描述
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                
                # 转换 Base64
                image_b64 = base64.b64encode(image_bytes).decode('utf-8')
                
                # 调用 Vision 模型描述图片
                print(f"Detecting image on page {page_num+1}, calling Vision model...")
                description = await describe_image(image_b64)
                if description:
                    text += f"\n[图片内容描述: {description}]\n"
                    
        doc.close()
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
    return text

def extract_text_from_docx(file_path: str) -> str:
    """从 Word 文档中提取文字，支持普通段落、表格、文本框"""
    text = ""
    try:
        doc = Document(file_path)

        # 1. 提取普通段落（包括各级标题）
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"

        # 2. 提取表格中的文字
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_texts.append(cell.text.strip())
                if row_texts:
                    text += " | ".join(row_texts) + "\n"

        # 3. 如果还是空的，尝试从 XML 中全量提取文字（针对文本框等特殊结构）
        if not text.strip():
            from lxml import etree
            ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            xml_bytes = doc.element.xml.encode('utf-8')
            root = etree.fromstring(xml_bytes)
            all_texts = root.iter(f'{{{ns}}}t')
            text = "\n".join(t.text for t in all_texts if t.text and t.text.strip())


    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
    return text


def clean_text(text: str) -> str:
    """清洗文本，去除零宽字符等不可见特殊字符"""
    # 去除各类零宽字符、控制字符等
    import re
    text = re.sub(r'[\u200b\u200c\u200d\ufeff\u00a0]', ' ', text)
    # 去除多余空行
    text = re.sub(r'\n\s*\n', '\n', text)
    return text.strip()

async def parse_document(file_path: str) -> str:
    """Parse document based on its extension (Async support for Vision model)."""
    import os
    import asyncio
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        text = await extract_text_from_pdf(file_path)
    elif ext in [".doc", ".docx"]:
        text = await asyncio.to_thread(extract_text_from_docx, file_path)
    elif ext in [".txt", ".md"]:
        def read_txt():
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        text = await asyncio.to_thread(read_txt)
    else:
        raise ValueError(f"Unsupported document format: {ext}")
    return clean_text(text)

def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """Split text into smaller chunks for vectorization."""
    # A simple character-based chunking; in production, token-based is better
    chunks = []
    start = 0
    text_length = len(text)
    while start < text_length:
        end = min(start + chunk_size, text_length)
        
        # If not at the end, try to find a natural break (like a newline or period)
        if end < text_length:
            # Look back for a period or newline within the last 50 characters
            lookback = text[end-50:end]
            if "\n" in lookback:
                end = end - 50 + lookback.rfind("\n") + 1
            elif "。" in lookback:
                end = end - 50 + lookback.rfind("。") + 1
            elif "." in lookback:
                end = end - 50 + lookback.rfind(".") + 1
                
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        if end >= text_length:
            break
            
        next_start = end - overlap
        if next_start <= start:
            next_start = start + 10 # Force advance if overlap causes infinite loop
        start = next_start
        
    return chunks
